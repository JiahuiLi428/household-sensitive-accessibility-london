from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[3]
SOURCE_DOCX = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8" / "newdraft_extension_integrated_fixed.docx"
OUT_DOCX = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8" / "newdraft_extension_integrated_final.docx"
FIGURE12_IMAGE = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8" / "assets" / "figure12_previous_locator.png"


FIGURE12_TEXT = (
    "Figure 12 locates the Barking Riverside / Thames View case within Greater London and East London before zooming to the local built environment. "
    "Panel A shows the case area's position within the metropolitan boundary; Panel B situates it within East London and the Thames-side urban edge; "
    "Panel C shows the selected LSOAs against roads, the River Thames and built-up blocks. The figure is used as a locator and spatial-context map for the case analysis."
)

FIGURE12_CAPTION = (
    "Figure 12. Location and local spatial context of the Barking Riverside / Thames View case area. "
    "Panel A locates the study area within Greater London; Panel B zooms to East London; Panel C shows the selected LSOAs in relation to the River Thames, roads and surrounding built-up fabric."
)

APPENDIX_D_SENTENCE = (
    "Appendix D further shows that the care-sensitive mismatch pattern remains visible under a finer residential-origin resolution, while revealing hidden within-LSOA pockets and centroid over- or under-estimation that the main LSOA model necessarily smooths."
)


def replace_paragraph_text(paragraph, text):
    paragraph._p.clear_content()
    paragraph.add_run(text)


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph.__class__(new_p, paragraph._parent)
    inserted.add_run(text)
    inserted.paragraph_format.space_after = Pt(6)
    return inserted


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def has_picture(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def find_paragraph(doc, needle):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Could not find paragraph containing: {needle}")


def replace_figure12(doc):
    replace_paragraph_text(
        find_paragraph(doc, "Figure 12 links the case comparison"),
        FIGURE12_TEXT,
    )

    caption = find_paragraph(doc, "Figure 12. Regional connectivity")
    paragraphs = doc.paragraphs
    caption_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for candidate in reversed(paragraphs[:caption_index]):
        if has_picture(candidate):
            image_para = candidate
            break
    else:
        raise ValueError("Could not find the image paragraph before Figure 12 caption")

    new_p = OxmlElement("w:p")
    image_para._p.addprevious(new_p)
    picture_para = caption.__class__(new_p, caption._parent)
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_para.add_run().add_picture(str(FIGURE12_IMAGE), width=Inches(6.4))
    delete_paragraph(image_para)

    replace_paragraph_text(caption, FIGURE12_CAPTION)


def move_appendix_d_sentence(doc):
    old = find_paragraph(doc, APPENDIX_D_SENTENCE)
    delete_paragraph(old)
    target = find_paragraph(doc, "This dissertation's contribution is to move from generic proximity")
    insert_paragraph_after(target, APPENDIX_D_SENTENCE)


def main():
    if not FIGURE12_IMAGE.exists():
        raise FileNotFoundError(FIGURE12_IMAGE)
    doc = Document(SOURCE_DOCX)
    replace_figure12(doc)
    move_appendix_d_sentence(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
