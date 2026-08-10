from pathlib import Path
import csv
import math

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[3]
SOURCE_DOCX = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8" / "newdraft.docx"
OUT_DOCX = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8" / "newdraft_extension_integrated_fixed.docx"
EXT = PROJECT_ROOT / "extension"
MAPS = EXT / "outputs" / "maps"
TABLES = EXT / "outputs" / "tables"
CASES = EXT / "outputs" / "case_studies"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=7.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fmt_value(value):
    if value is None:
        return ""
    text = str(value).strip()
    if text == "" or text.lower() == "nan":
        return ""
    try:
        number = float(text)
        if math.isfinite(number):
            if abs(number - round(number)) < 1e-9:
                return str(int(round(number)))
            return f"{number:.3f}".rstrip("0").rstrip(".")
    except ValueError:
        pass
    return text


def read_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def replace_paragraph_text(paragraph, text):
    paragraph._p.clear_content()
    paragraph.add_run(text)


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph.__class__(new_p, paragraph._parent)
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    inserted.paragraph_format.space_after = Pt(6)
    return inserted


def find_first_paragraph(doc, needle):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Could not find paragraph containing: {needle}")


def replace_text_everywhere(doc, old, new):
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            replace_paragraph_text(paragraph, paragraph.text.replace(old, new))


def apply_main_text_fixes(doc):
    replace_paragraph_text(
        find_first_paragraph(doc, "Global Moran's I for caregiving adult mismatch is 0.679"),
        "The LISA analysis, shown in Figure 7D, adds an explicitly spatial test. Global Moran's I for caregiving adult mismatch is 0.682 (z = 81.39, p < 0.001), indicating strong positive spatial autocorrelation, and Local Moran's I identifies 670 LSOAs - 13.42% of London - as high-high mismatch clusters. These concentrate in outer boroughs including Hillingdon, Bromley, Havering, Enfield, Bexley, Redbridge, Harrow, Croydon and Barking and Dagenham. Across the high-high cluster, mean caregiving adult mismatch is 1.80, mean caregiving demand z-score is 1.09 and mean caregiving access is 24.72, reinforcing the claim that the most spatially clustered care-access deficits are not simply an inner-city deprivation story."
    )
    replace_text_everywhere(doc, "1,185 LSOAs (23.73% of London)", "670 LSOAs (13.42% of London)")
    replace_text_everywhere(doc, "1,185 LISA high-high LSOAs", "670 LISA high-high LSOAs")
    replace_text_everywhere(doc, "1,185 LSOAs as high-high mismatch clusters", "670 LSOAs as high-high mismatch clusters")
    insert_paragraph_after(
        find_first_paragraph(doc, "Waiting time and access/egress walking are handled internally"),
        "Because representative LSOA origin points can understate internal residential variation, Appendix D adds a supplementary 100m residential-grid robustness analysis. The extension keeps Census need at LSOA level but recalculates the accessibility component from residential grid origins, then compares centroid-based and grid-informed results."
    )
    note = insert_paragraph_after(
        find_first_paragraph(doc, "Table 6. Facility-specific equity ranking under the walking scenario."),
        "Note: Table 6 reports effect sizes for ranking completeness. The hospital effect size is very small and Table 5 shows the hospital result is not statistically significant, so its rank should not be read as evidence of a meaningful deprivation gradient."
    )
    note.runs[0].italic = True
    insert_paragraph_after(
        find_first_paragraph(doc, "Future research should refine this boundary along three lines"),
        "The supplementary grid analysis in Appendix D partially addresses the centroid-origin limitation by repeating the accessibility component on a 100m residential grid. It is a robustness check on accessibility precision, not a redistribution of Census need to individual origins."
    )
    insert_paragraph_after(
        find_first_paragraph(doc, "The central implication, finally, is that London's 15-minute city"),
        "Appendix D further shows that the care-sensitive mismatch pattern remains visible under a finer residential-origin resolution, while revealing hidden within-LSOA pockets and centroid over- or under-estimation that the main LSOA model necessarily smooths."
    )


def add_note(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(text)


def add_caption(doc, caption):
    p = doc.add_paragraph()
    p.style = "Caption" if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)


def add_picture(doc, image_path, caption, width=6.3):
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_table_from_rows(doc, rows, columns, caption, max_rows=None, font_size=7.2):
    doc.add_paragraph(caption, style="Caption" if "Caption" in [s.name for s in doc.styles] else None)
    use_rows = rows if max_rows is None else rows[:max_rows]
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for j, col in enumerate(columns):
        set_cell_text(hdr[j], col, bold=True, size=font_size)
        set_cell_shading(hdr[j], "D9EAF7")
    for row in use_rows:
        cells = table.add_row().cells
        for j, col in enumerate(columns):
            set_cell_text(cells[j], fmt_value(row.get(col, "")), size=font_size)
    if max_rows is not None and len(rows) > max_rows:
        p = doc.add_paragraph()
        p.add_run(f"Note: table truncated to first {max_rows} rows in the manuscript appendix; the full borough table should be supplied as supplementary material or deposited with the project repository.").italic = True
        p.paragraph_format.space_after = Pt(6)
    else:
        doc.add_paragraph()
    return table


def add_case_tables(doc, rows):
    metric_rows = []
    interpretation_rows = []
    for row in rows:
        short = row["case_name"].replace("Case 1 ", "Case 1: ").replace("Case 2 ", "Case 2: ").replace("Case 3 ", "Case 3: ")
        metric_rows.append({
            "case": short,
            "care_need": row.get("care_need", ""),
            "child_need": row.get("child_need", ""),
            "older_need": row.get("older_need", ""),
            "centroid_access": row.get("centroid_access", ""),
            "grid_mean_access": row.get("grid_mean_access", ""),
            "lowest_decile_access": row.get("lowest_decile_access", ""),
            "within_lsoa_variation": row.get("within_lsoa_variation", ""),
            "mismatch_z": row.get("mismatch_z", row.get("mismatch_score", "")),
        })
        interpretation_rows.append({
            "case": short,
            "case_type": row.get("case_type", ""),
            "lisa_class": row.get("lisa_class", ""),
            "compound_type": row.get("compound_type", ""),
            "main_service_deficit": row.get("main_service_deficit", ""),
            "policy_implication": row.get("policy_implication", ""),
        })
    add_table_from_rows(
        doc,
        metric_rows,
        ["case", "care_need", "child_need", "older_need", "centroid_access", "grid_mean_access", "lowest_decile_access", "within_lsoa_variation", "mismatch_z"],
        "Table D5a. Case-study comparison: need, accessibility and z-standardised grid-informed mismatch metrics.",
        font_size=6.7,
    )
    doc.add_page_break()
    add_table_from_rows(
        doc,
        interpretation_rows,
        ["case", "case_type", "lisa_class", "compound_type", "main_service_deficit", "policy_implication"],
        "Table D5b. Case-study comparison: interpretation and policy implication.",
        font_size=7.0,
    )


def add_markdown_paragraphs(doc, path):
    for line in path.read_text(encoding="utf-8").splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("# "):
            continue
        if text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text[2:])
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)


def main():
    doc = Document(SOURCE_DOCX)
    apply_main_text_fixes(doc)

    # Keep the original document intact; append the extension as Appendix D.
    doc.add_page_break()
    doc.add_heading("Appendix D. Multi-scale robustness and comparative neighbourhood analysis", level=1)
    add_note(
        doc,
        "Purpose.",
        "This appendix integrates the extension analysis requested after the main dissertation draft: a 100m residential-grid robustness analysis, centroid-grid comparison, grid-informed mismatch and LISA, compound mismatch, and three comparative neighbourhood cases."
    )

    doc.add_heading("D.1 Extension workflow and methodological logic", level=2)
    add_markdown_paragraphs(doc, EXT / "extension_methodology.md")
    add_note(
        doc,
        "Scale note.",
        "Section 4.10 reports Barking Riverside / Thames View using raw need-access gap values, whereas Appendix D reports z-standardised mismatch scores for LISA comparability. The two numbers are therefore different scales, not conflicting estimates."
    )
    add_picture(doc, MAPS / "figure_D1_extension_workflow.png", "Figure D1. Extension workflow diagram.", width=6.4)

    doc.add_heading("D.2 Grid construction and residential filtering", level=2)
    add_table_from_rows(
        doc,
        read_csv(TABLES / "table_D1_grid_construction_summary.csv"),
        ["metric", "value"],
        "Table D1. Grid construction summary.",
        font_size=8.2,
    )

    doc.add_heading("D.3 Centroid-grid comparison", level=2)
    add_note(
        doc,
        "Interpretation.",
        "The grid-informed results preserve the original LSOA need framework but replace the single origin point with residential 100m grid origins. The comparison therefore tests robustness of accessibility estimates, not a new population-need model."
    )
    add_picture(doc, MAPS / "figure_D2_centroid_grid_accessibility.png", "Figure D2. Centroid-based and grid-informed accessibility maps.", width=6.4)
    add_picture(doc, MAPS / "figure_D3_centroid_minus_grid_difference.png", "Figure D3. Centroid minus grid-informed accessibility.", width=6.2)
    add_picture(doc, MAPS / "figure_D4_within_lsoa_variation.png", "Figure D4. Within-LSOA accessibility variation.", width=6.2)
    stats_rows = read_csv(TABLES / "table_D2_centroid_grid_comparison_statistics.csv")
    add_table_from_rows(
        doc,
        stats_rows,
        ["group", "mode", "pearson_correlation", "spearman_rank_correlation", "mean_absolute_difference", "rmse", "lowest_accessibility_20_jaccard", "highest_mismatch_20_jaccard"],
        "Table D2. Centroid-grid comparison statistics.",
        font_size=6.6,
    )

    doc.add_heading("D.4 Grid-informed mismatch, LISA and compound mismatch", level=2)
    add_picture(doc, MAPS / "figure_D5_lisa_reclassification.png", "Figure D5. LISA high-high reclassification.", width=6.2)
    add_picture(doc, MAPS / "figure_D6_compound_mismatch.png", "Figure D6. Compound mismatch map.", width=6.2)
    add_table_from_rows(
        doc,
        read_csv(TABLES / "table_D3_lisa_reclassification.csv"),
        ["group", "comparison", "lisa_reclassification", "N", "share_lsoas_pct"],
        "Table D3. LISA reclassification and cluster summary.",
        font_size=7.0,
    )
    add_table_from_rows(
        doc,
        read_csv(TABLES / "table_D4_compound_mismatch_by_borough.csv"),
        ["borough", "borough_lsoa_count", "single_group_count", "two_group_count", "three_group_count", "share_of_borough_lsoas"],
        "Table D4. Compound mismatch by borough.",
        max_rows=20,
        font_size=6.8,
    )

    doc.add_heading("D.5 Comparative neighbourhood cases", level=2)
    add_note(
        doc,
        "Case-selection rule.",
        "Case 1 is fixed as Barking Riverside / Thames View; Case 2 is selected from compound high-high mismatch candidates; Case 3 is a positive counter-case matched on need-related variables but with higher grid-informed accessibility and lower mismatch."
    )
    add_note(
        doc,
        "Counter-case note.",
        "The Newham counter-case in Section 4.10 remains the original LSOA-scale narrative comparator. The Barking and Dagenham 015F case here is a separate, algorithmically selected grid-robustness counter-case matched on need variables; it tests local robustness and does not replace the Newham comparison."
    )
    add_picture(doc, MAPS / "figure_D7_case_study_comparison.png", "Figure D7. Three case-study comparison.", width=6.4)
    for slug, title in [
        ("case_1_barking_riverside_thames_view", "Figure D8. Case 1: Barking Riverside / Thames View."),
        ("case_2_compound_mismatch", "Figure D9. Case 2: Hillingdon 003D compound mismatch case."),
        ("case_3_positive_counter_case", "Figure D10. Case 3: Barking and Dagenham 015F positive counter-case."),
    ]:
        add_picture(doc, CASES / f"{slug}_map.png", title, width=6.1)
    add_case_tables(doc, read_csv(TABLES / "table_D5_case_study_comparison.csv"))

    doc.add_heading("D.6 Results summary and limitations", level=2)
    add_markdown_paragraphs(doc, EXT / "extension_results_summary.md")
    doc.add_heading("D.7 Limitations of the extension", level=2)
    add_markdown_paragraphs(doc, EXT / "extension_limitations.md")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
