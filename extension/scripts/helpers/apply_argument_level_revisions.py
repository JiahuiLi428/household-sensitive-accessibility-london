from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[3]
DOC_DIR = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8"
SOURCE_DOCX = DOC_DIR / "newdraft_extension_integrated_final.docx"
FIXED_DOCX = DOC_DIR / "newdraft_extension_integrated_fixed.docx"
OUT_DOCX = DOC_DIR / "newdraft_extension_integrated_distinction.docx"
ASSET_DIR = DOC_DIR / "assets"
FIGURE12_ANALYTIC = ASSET_DIR / "figure12_analytic_care_chain.png"


FIGURE12_TEXT = (
    "Figure 12 links the case comparison to the dissertation's care-chain argument. Panel A locates the selected high-high mismatch cluster within London; "
    "Panel B shows the local Barking Riverside / Thames View service environment, including illustrative adult and caregiver walking-distance equivalents, facilities and transport entry points; "
    "and Panel C summarises why regional connectivity does not automatically solve local care-chain accessibility. This figure therefore does more than orient the reader: it visualises the mechanism behind the accessibility paradox."
)

FIGURE12_CAPTION = (
    "Figure 12. Regional connectivity and local care-chain mismatch in the Barking Riverside / Thames View case. Panel A locates the selected high-high care-mismatch cluster within London. "
    "Panel B maps the local everyday-service environment and illustrative adult/caregiver walking-distance equivalents. Panel C summarises the care-chain logic: accessibility depends on combining services within constrained daily routines, not on single-destination proximity alone."
)


def replace_paragraph_text(paragraph, text):
    paragraph._p.clear_content()
    paragraph.add_run(text)


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph.__class__(new_p, paragraph._parent)
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    inserted.paragraph_format.space_after = Pt(6)
    return inserted


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def has_picture(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def find_paragraph(doc, needle):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Could not find paragraph containing: {needle}")


def find_previous_picture_paragraph(doc, caption):
    paragraphs = doc.paragraphs
    caption_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for candidate in reversed(paragraphs[:caption_index]):
        if has_picture(candidate):
            return candidate
    raise ValueError("Could not find picture paragraph before caption")


def extract_analytic_figure12():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if FIGURE12_ANALYTIC.exists():
        return FIGURE12_ANALYTIC
    fixed = Document(FIXED_DOCX)
    caption = find_paragraph(fixed, "Figure 12. Regional connectivity")
    picture_para = find_previous_picture_paragraph(fixed, caption)
    blips = picture_para._p.xpath(".//a:blip")
    if not blips:
        raise ValueError("Figure 12 picture paragraph did not contain an embedded image")
    rid = blips[0].get(qn("r:embed"))
    image_part = fixed.part.related_parts[rid]
    FIGURE12_ANALYTIC.write_bytes(image_part.blob)
    return FIGURE12_ANALYTIC


def replace_figure12(doc):
    image_path = extract_analytic_figure12()
    replace_paragraph_text(
        find_paragraph(doc, "Figure 12 locates the Barking Riverside"),
        FIGURE12_TEXT,
    )
    caption = find_paragraph(doc, "Figure 12. Location and local spatial context")
    old_picture = find_previous_picture_paragraph(doc, caption)
    new_p = OxmlElement("w:p")
    old_picture._p.addprevious(new_p)
    picture_para = caption.__class__(new_p, caption._parent)
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_para.add_run().add_picture(str(image_path), width=Inches(6.4))
    delete_paragraph(old_picture)
    replace_paragraph_text(caption, FIGURE12_CAPTION)


def add_argument_level_revisions(doc):
    insert_paragraph_after(
        find_paragraph(doc, "Recent accessibility research strengthens this operational case further"),
        "The distinction from Soukhov et al. (2025) is important. Their contribution is to show how mobility-of-care destinations can be brought into accessibility measurement; this dissertation extends that logic in a different direction by combining role-specific service baskets with Census-derived need, IMD context, signed mismatch and LISA clustering across all London LSOAs. In other words, care is not only added as another destination category: it becomes the basis for testing whether local service environments match household responsibility.",
    )
    insert_paragraph_after(
        find_paragraph(doc, "Critical and behavioural perspectives add a caution"),
        "Caprotti, Duarte and Joss's (2024) critique of the 15-minute city as 'paranoid urbanism' is especially useful here because it warns against treating proximity as a politically neutral good. The dissertation does not reject the 15-minute-city idea, but it treats the concept as conditional: proximity becomes equitable only if it is evaluated against who performs care, who depends on local services, and whether theoretical reachability can be converted into usable everyday access.",
    )
    insert_paragraph_after(
        find_paragraph(doc, "The three accessibility subjects are deliberately asymmetric"),
        "This is also a gender-aware but not gender-disaggregated study. The mobility-of-care literature shows that escorting, shopping, healthcare coordination and household maintenance are often gendered, but the Census and routing data used here do not identify who performs these tasks within each household. The caregiving-adult category should therefore be read as a spatial proxy for care responsibility and service-bundle pressure, not as a direct estimate of women's mobility or individual caregiver behaviour.",
    )
    insert_paragraph_after(
        find_paragraph(doc, "Because these choices remain contestable"),
        "Reflexively, the weights should be understood as an explicit modelling judgement rather than a discovered fact. Their purpose is not to claim that, for example, a pharmacy is universally worth a fixed fraction of a caregiver's daily life; it is to make the normative assumptions of a household-sensitive accessibility test visible and contestable. The sensitivity analysis therefore functions as an accountability check: if the core geography disappeared under plausible alternative weights, the argument would be much weaker.",
    )

    result_heading = find_paragraph(doc, "4. Results: Research Questions and Evidence Sets")
    heading = OxmlElement("w:p")
    result_heading._p.addprevious(heading)
    heading_para = result_heading.__class__(heading, result_heading._parent)
    heading_para.style = "Heading 2"
    heading_para.add_run("3.7 Validity, reliability and robustness")
    insert_paragraph_after(
        heading_para,
        "Construct validity is addressed by making each conceptual claim correspond to an explicit indicator: accessibility is measured through routable travel time and facility scores; need is measured through Census age and household-composition proxies; and mismatch is measured as the signed gap between standardised need and accessibility. This does not make the indicators perfect, but it keeps the chain from concept to measure visible.",
    )
    insert_paragraph_after(
        find_paragraph(doc, "Construct validity is addressed"),
        "Reliability is supported by using the same LSOA boundaries, facility categories, routing assumptions, basket weights and standardisation logic across all groups and scenarios. The analysis is scripted rather than manually classified, random-weight sensitivity checks use a fixed seed, and Appendix D tests whether the main accessibility and mismatch patterns survive a finer 100m residential-origin grid.",
    )
    insert_paragraph_after(
        find_paragraph(doc, "Reliability is supported"),
        "External validity is necessarily more modest. The results should not be read as individual travel behaviour or as a universal model of care in every city. Their stronger claim is analytic transferability: in any city where average proximity is used as an equity claim, the same question should be asked about household need, bundled services and practical usability.",
    )

    insert_paragraph_after(
        find_paragraph(doc, "The planning implication is practical, not just conceptual"),
        "In London policy terms, this aligns directly with the Healthy Streets Approach and the Mayor's Transport Strategy, both of which frame walking, public transport and street quality as health and equity interventions rather than only movement infrastructure (Transport for London, 2017; Greater London Authority, 2018). The dissertation's contribution is to add a care-sensitive diagnostic layer to those frameworks: Healthy Streets indicators can improve the micro-environment, but they should be targeted where service-basket mismatch shows that everyday care chains are hardest to complete.",
    )

    delete_paragraph(find_paragraph(doc, "Throughout this section, higher IMD deciles indicate less deprived LSOAs."))


def add_references(doc):
    if any("Greater London Authority (2018) Mayor's Transport Strategy" in p.text for p in doc.paragraphs):
        return
    geurs = find_paragraph(doc, "Geurs, K. T. and van Wee")
    insert_paragraph_after(
        geurs,
        "Greater London Authority (2018) Mayor's Transport Strategy 2018. Available at: https://www.london.gov.uk/programmes-strategies/transport/our-vision-transport/mayors-transport-strategy-2018 (Accessed: 17 July 2026).",
    )
    tfl_2026 = find_paragraph(doc, "Transport for London (2026a)")
    new_ref = OxmlElement("w:p")
    tfl_2026._p.addprevious(new_ref)
    ref_para = tfl_2026.__class__(new_ref, tfl_2026._parent)
    ref_para.add_run("Transport for London (2017) Healthy Streets for London. Available at: https://content.tfl.gov.uk/healthy-streets-for-london.pdf (Accessed: 17 July 2026).")
    ref_para.paragraph_format.space_after = Pt(6)


def main():
    doc = Document(SOURCE_DOCX)
    replace_figure12(doc)
    add_argument_level_revisions(doc)
    add_references(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
