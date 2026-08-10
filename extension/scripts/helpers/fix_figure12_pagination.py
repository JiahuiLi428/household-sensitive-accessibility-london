from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[3]
DOC_DIR = PROJECT_ROOT / "outputs" / "manuscript" / "current" / "7.8"
SOURCE_DOCX = DOC_DIR / "newdraft_extension_integrated_distinction.docx"
OUT_DOCX = DOC_DIR / "newdraft_extension_integrated_distinction_pagination_fixed.docx"


def has_picture(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def find_paragraph(doc, needle):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Could not find paragraph containing: {needle}")


def find_previous_picture_paragraph(doc, caption):
    paragraphs = doc.paragraphs
    caption_index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    for candidate in reversed(paragraphs[:caption_index]):
        if has_picture(candidate):
            return candidate
    raise ValueError("Could not find picture paragraph before Figure 12 caption")


def main():
    doc = Document(SOURCE_DOCX)

    intro = find_paragraph(doc, "Figure 12 links the case comparison")
    caption = find_paragraph(doc, "Figure 12. Regional connectivity and local care-chain mismatch")
    picture = find_previous_picture_paragraph(doc, caption)

    # Keep the Figure 12 setup paragraph, image, and caption as one pagination unit.
    intro.paragraph_format.keep_with_next = True
    intro.paragraph_format.keep_together = True
    picture.paragraph_format.keep_with_next = True
    picture.paragraph_format.keep_together = True
    caption.paragraph_format.keep_together = True

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
